import streamlit as st
from openai import OpenAI, RateLimitError, AuthenticationError
from pypdf import PdfReader
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import pandas as pd
import requests
import itertools
import re
import os
import json
from datetime import datetime


# =====================================================
# CONFIGURACION GENERAL
# =====================================================

st.set_page_config(
    page_title="Asistente Cientifico de Articulos",
    layout="wide"
)

st.title("📚 Asistente Científico para Elaboración de Artículos Académicos")

st.info(
    "Aplicación para apoyar la elaboración de artículos científicos a partir de PDF, "
    "Zotero, matrices bibliográficas, redes bibliométricas para VOSviewer, revisión "
    "académica, revisión de coherencia y exportación en Word, PDF, Excel y TXT."
)


# =====================================================
# API OPENAI
# =====================================================

try:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
except Exception:
    openai_api_key = None

if not openai_api_key:
    st.error(
        "No se encontró OPENAI_API_KEY. Cree el archivo .streamlit/secrets.toml "
        "con la clave de OpenAI."
    )
    st.stop()

client = OpenAI(api_key=openai_api_key)


# =====================================================
# ESTADO DE SESION
# =====================================================

estado_inicial = {
    "fuentes_extraidas": "",
    "metadatos_zotero": [],
    "matriz_bibliografica": "",
    "referencias_extraidas": "",
    "articulo_completo": "",
    "evaluacion_calidad": "",
    "revision_coherencia": "",
    "resumen_abstract": "",
    "parafrasis": "",
    "map_keywords": "",
    "network_keywords": "",
    "map_authors": "",
    "network_authors": "",
    "proyecto_actual": "",
}

for clave, valor in estado_inicial.items():
    if clave not in st.session_state:
        st.session_state[clave] = valor

if "secciones_generadas" not in st.session_state:
    st.session_state.secciones_generadas = {}

if "datos_articulo" not in st.session_state:
    st.session_state.datos_articulo = {}


# =====================================================
# FUNCIONES OPENAI
# =====================================================

def llamar_openai(prompt, modelo="gpt-4o-mini", temperatura=0.3):
    try:
        respuesta = client.chat.completions.create(
            model=modelo,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un asistente académico experto en investigación científica, "
                        "redacción de artículos, revisión bibliográfica, metodología, "
                        "bibliometría y normas de citación. No debes inventar autores, "
                        "años, DOI, revistas, datos ni resultados."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=temperatura
        )

        return respuesta.choices[0].message.content

    except AuthenticationError:
        return "ERROR: La clave API de OpenAI no es válida."

    except RateLimitError:
        return (
            "ERROR: La cuenta de OpenAI no tiene cuota disponible o superó el límite. "
            "Revise Billing y Usage Limits en OpenAI Platform."
        )

    except Exception as e:
        return f"ERROR: {e}"


# =====================================================
# FUNCIONES PDF
# =====================================================

def extraer_texto_pdf(archivo_pdf):
    texto = ""

    try:
        lector = PdfReader(archivo_pdf)

        for numero_pagina, pagina in enumerate(lector.pages, start=1):
            contenido = pagina.extract_text()

            if contenido:
                texto += f"\n\n--- Página {numero_pagina} ---\n"
                texto += contenido

    except Exception as e:
        texto = f"Error al leer PDF: {e}"

    return texto


def recortar_texto(texto, max_caracteres):
    if len(texto) > max_caracteres:
        return (
            texto[:max_caracteres]
            + "\n\n[Texto recortado automáticamente por longitud.]"
        )

    return texto


def limpiar_texto_pdf(texto):
    reemplazos = {
        "📚": "",
        "📄": "",
        "✅": "",
        "✓": "",
        "✔": "",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "–": "-",
        "—": "-",
        "•": "-",
        "…": "...",
        "α": "alpha",
        "β": "beta",
        "γ": "gamma",
    }

    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)

    return texto


def generar_pdf(texto):
    texto = limpiar_texto_pdf(texto)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    for linea in texto.split("\n"):
        if linea.strip() == "":
            pdf.ln(4)
        else:
            try:
                pdf.multi_cell(
                    0,
                    8,
                    linea.encode("latin-1", errors="ignore").decode("latin-1")
                )
            except Exception:
                pdf.multi_cell(0, 8, "[Texto omitido por caracteres incompatibles]")

    return pdf.output(dest="S").encode("latin-1", errors="ignore")


def generar_docx(texto, titulo_documento="Articulo cientifico generado"):
    documento = Document()

    section = documento.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    documento.styles["Normal"].font.name = "Times New Roman"
    documento.styles["Normal"].font.size = Pt(12)

    documento.add_heading(titulo_documento, level=0)

    encabezados = [
        "resumen",
        "abstract",
        "palabras clave",
        "keywords",
        "introducción",
        "introduccion",
        "revisión de literatura",
        "revision de literatura",
        "marco teórico",
        "marco teorico",
        "metodología",
        "metodologia",
        "resultados",
        "discusión",
        "discusion",
        "conclusiones",
        "limitaciones",
        "referencias"
    ]

    for linea in texto.split("\n"):
        linea_limpia = linea.strip()

        if linea_limpia == "":
            documento.add_paragraph("")
        elif linea_limpia.lower() in encabezados:
            documento.add_heading(linea_limpia, level=1)
        elif re.match(r"^\d+[\.\)]\s+", linea_limpia):
            documento.add_heading(linea_limpia, level=1)
        else:
            documento.add_paragraph(linea_limpia)

    archivo = BytesIO()
    documento.save(archivo)
    archivo.seek(0)
    return archivo


def generar_excel_desde_texto(texto, hoja="Matriz"):
    filas = []

    for linea in texto.split("\n"):
        if linea.strip():
            filas.append({"Contenido": linea.strip()})

    df = pd.DataFrame(filas)

    archivo = BytesIO()
    with pd.ExcelWriter(archivo, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=hoja)

    archivo.seek(0)
    return archivo


# =====================================================
# FUNCIONES ZOTERO
# =====================================================

def zotero_base_url(library_type, library_id):
    if library_type == "user":
        return f"https://api.zotero.org/users/{library_id}"
    return f"https://api.zotero.org/groups/{library_id}"


def zotero_headers(api_key):
    return {
        "Zotero-API-Key": api_key,
        "Zotero-API-Version": "3"
    }


def zotero_get_paginated(url, api_key, params=None):
    if params is None:
        params = {}

    params["limit"] = 100
    resultados = []
    siguiente_url = url

    while siguiente_url:
        respuesta = requests.get(
            siguiente_url,
            headers=zotero_headers(api_key),
            params=params if siguiente_url == url else None,
            timeout=60
        )

        if respuesta.status_code != 200:
            raise Exception(
                f"Error Zotero {respuesta.status_code}: {respuesta.text[:500]}"
            )

        resultados.extend(respuesta.json())

        link = respuesta.headers.get("Link", "")
        siguiente_url = None

        partes = link.split(",")

        for parte in partes:
            if 'rel="next"' in parte:
                match = re.search(r"<([^>]+)>", parte)
                if match:
                    siguiente_url = match.group(1)
                    break

        params = {}

    return resultados


def zotero_listar_colecciones(library_type, library_id, api_key):
    base = zotero_base_url(library_type, library_id)
    url = f"{base}/collections"

    colecciones = zotero_get_paginated(url, api_key)

    datos = []

    for c in colecciones:
        data = c.get("data", {})
        datos.append({
            "key": data.get("key", c.get("key", "")),
            "name": data.get("name", "Sin nombre"),
            "numItems": data.get("numItems", 0)
        })

    return datos


def zotero_items_coleccion(library_type, library_id, api_key, collection_key):
    base = zotero_base_url(library_type, library_id)
    url = f"{base}/collections/{collection_key}/items"

    params = {
        "itemType": "-attachment"
    }

    return zotero_get_paginated(url, api_key, params=params)


def zotero_children(library_type, library_id, api_key, item_key):
    base = zotero_base_url(library_type, library_id)
    url = f"{base}/items/{item_key}/children"

    return zotero_get_paginated(url, api_key)


def zotero_descargar_pdf_adjunto(library_type, library_id, api_key, attachment_key):
    base = zotero_base_url(library_type, library_id)
    url = f"{base}/items/{attachment_key}/file"

    respuesta = requests.get(
        url,
        headers=zotero_headers(api_key),
        timeout=120,
        allow_redirects=True
    )

    if respuesta.status_code == 200 and respuesta.content:
        return BytesIO(respuesta.content)

    return None


def autores_zotero(creators):
    autores = []

    for creador in creators:
        if creador.get("creatorType") in ["author", "editor"]:
            nombre = creador.get("name")
            if not nombre:
                first = creador.get("firstName", "")
                last = creador.get("lastName", "")
                nombre = f"{first} {last}".strip()

            if nombre:
                autores.append(nombre)

    return autores


def item_zotero_a_metadato(item):
    data = item.get("data", {})

    autores = autores_zotero(data.get("creators", []))
    tags = [t.get("tag", "") for t in data.get("tags", []) if t.get("tag")]

    metadato = {
        "key": data.get("key", item.get("key", "")),
        "itemType": data.get("itemType", ""),
        "title": data.get("title", ""),
        "authors": autores,
        "year": data.get("date", "")[:4],
        "date": data.get("date", ""),
        "publicationTitle": data.get("publicationTitle", ""),
        "journalAbbreviation": data.get("journalAbbreviation", ""),
        "DOI": data.get("DOI", ""),
        "url": data.get("url", ""),
        "abstractNote": data.get("abstractNote", ""),
        "tags": tags,
        "language": data.get("language", ""),
        "volume": data.get("volume", ""),
        "issue": data.get("issue", ""),
        "pages": data.get("pages", "")
    }

    return metadato


def metadatos_a_texto(metadatos):
    texto = ""

    for i, m in enumerate(metadatos, start=1):
        texto += f"\n\n--- Registro Zotero {i} ---\n"
        texto += f"Título: {m.get('title', '')}\n"
        texto += f"Autores: {', '.join(m.get('authors', []))}\n"
        texto += f"Año: {m.get('year', '')}\n"
        texto += f"Revista/Fuente: {m.get('publicationTitle', '')}\n"
        texto += f"DOI: {m.get('DOI', '')}\n"
        texto += f"URL: {m.get('url', '')}\n"
        texto += f"Palabras clave / etiquetas: {', '.join(m.get('tags', []))}\n"
        texto += f"Resumen: {m.get('abstractNote', '')}\n"

    return texto


def metadatos_a_dataframe(metadatos):
    filas = []

    for m in metadatos:
        filas.append({
            "Titulo": m.get("title", ""),
            "Autores": "; ".join(m.get("authors", [])),
            "Año": m.get("year", ""),
            "Fuente": m.get("publicationTitle", ""),
            "DOI": m.get("DOI", ""),
            "URL": m.get("url", ""),
            "Palabras clave": "; ".join(m.get("tags", [])),
            "Resumen": m.get("abstractNote", "")
        })

    return pd.DataFrame(filas)


# =====================================================
# FUNCIONES VOSVIEWER
# =====================================================

def normalizar_termino(texto):
    texto = texto.lower().strip()
    texto = re.sub(r"\s+", " ", texto)
    texto = re.sub(r"[^\w\sáéíóúñü-]", "", texto)
    return texto


def extraer_keywords_metadatos(metadatos):
    documentos_keywords = []

    for m in metadatos:
        kws = []

        for tag in m.get("tags", []):
            tag_norm = normalizar_termino(tag)
            if len(tag_norm) > 2:
                kws.append(tag_norm)

        if not kws:
            titulo = m.get("title", "")
            palabras = re.findall(r"\b[a-zA-ZáéíóúÁÉÍÓÚñÑ]{5,}\b", titulo.lower())
            stopwords = {
                "using", "with", "from", "that", "this", "these", "those",
                "para", "desde", "sobre", "entre", "como", "este", "esta",
                "estos", "estas", "the", "and", "for", "del", "las", "los"
            }
            kws = [normalizar_termino(p) for p in palabras if p not in stopwords]

        documentos_keywords.append(sorted(set(kws)))

    return documentos_keywords


def extraer_autores_metadatos(metadatos):
    documentos_autores = []

    for m in metadatos:
        autores = [normalizar_termino(a) for a in m.get("authors", []) if a.strip()]
        documentos_autores.append(sorted(set(autores)))

    return documentos_autores


def construir_red_coocurrencia(documentos_items, min_ocurrencias=1):
    ocurrencias = {}
    enlaces = {}

    for items in documentos_items:
        items_unicos = sorted(set([i for i in items if i]))

        for item in items_unicos:
            ocurrencias[item] = ocurrencias.get(item, 0) + 1

        for a, b in itertools.combinations(items_unicos, 2):
            par = tuple(sorted([a, b]))
            enlaces[par] = enlaces.get(par, 0) + 1

    items_filtrados = {
        item: count
        for item, count in ocurrencias.items()
        if count >= min_ocurrencias
    }

    ids = {
        item: idx + 1
        for idx, item in enumerate(sorted(items_filtrados.keys()))
    }

    enlaces_filtrados = {
        par: strength
        for par, strength in enlaces.items()
        if par[0] in ids and par[1] in ids
    }

    return ids, items_filtrados, enlaces_filtrados


def generar_archivos_vosviewer(ids, ocurrencias, enlaces):
    map_lines = ["id\tlabel\tweight<Occurrences>"]

    for item, id_item in ids.items():
        map_lines.append(f"{id_item}\t{item}\t{ocurrencias.get(item, 0)}")

    network_lines = ["source\ttarget\tstrength"]

    for (a, b), strength in enlaces.items():
        network_lines.append(f"{ids[a]}\t{ids[b]}\t{strength}")

    return "\n".join(map_lines), "\n".join(network_lines)


def generar_red_keywords_vosviewer(metadatos, min_ocurrencias=1):
    docs = extraer_keywords_metadatos(metadatos)
    ids, ocurrencias, enlaces = construir_red_coocurrencia(docs, min_ocurrencias)
    return generar_archivos_vosviewer(ids, ocurrencias, enlaces)


def generar_red_autores_vosviewer(metadatos, min_ocurrencias=1):
    docs = extraer_autores_metadatos(metadatos)
    ids, ocurrencias, enlaces = construir_red_coocurrencia(docs, min_ocurrencias)
    return generar_archivos_vosviewer(ids, ocurrencias, enlaces)


# =====================================================
# PLANTILLAS
# =====================================================

def obtener_plantilla(nombre):
    plantillas = {
        "Artículo empírico tipo revista indexada": """
1. Título
2. Resumen
   - Problema
   - Objetivo
   - Metodología
   - Resultados principales
   - Conclusión central
3. Palabras clave
4. Introducción
   - Contextualización del problema
   - Justificación
   - Brecha investigativa
   - Objetivo
   - Organización del artículo
5. Revisión de literatura
   - Antecedentes
   - Enfoques teóricos
   - Estudios relacionados
   - Vacíos identificados
6. Metodología
   - Enfoque
   - Tipo de estudio
   - Fuentes o datos utilizados
   - Procedimiento
   - Técnicas de análisis
7. Resultados
8. Discusión
9. Conclusiones
10. Limitaciones y futuras investigaciones
11. Referencias
""",

        "Artículo de revisión": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Metodología de revisión
6. Desarrollo temático
7. Tendencias de la literatura
8. Discusión
9. Conclusiones
10. Referencias
""",

        "Artículo IMRyD": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Metodología
6. Resultados
7. Discusión
8. Conclusiones
9. Referencias
""",

        "Artículo de reflexión": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Planteamiento del problema
6. Desarrollo argumentativo
7. Discusión crítica
8. Conclusiones
9. Referencias
"""
    }

    return plantillas.get(nombre, plantillas["Artículo empírico tipo revista indexada"])


# =====================================================
# PROMPTS
# =====================================================

def prompt_matriz_bibliografica(fuentes, metadatos_texto, estilo_citacion):
    return f"""
Elabora una matriz bibliográfica rigurosa en español.

Usa los metadatos de Zotero y los textos fuente cuando estén disponibles.

Estructura de la matriz:

| Autor(es) | Año | Título | Objetivo | Metodología | Datos/Muestra | Resultados principales | Limitaciones | Aporte al tema | Referencia en {estilo_citacion} |

REGLAS:
- No inventes autores, años, DOI, revistas, resultados ni datos.
- Si un dato no aparece, escribe: No identificado.
- Si una referencia está incompleta, escribe: Referencia incompleta.
- Usa lenguaje académico.
- Prioriza precisión sobre extensión.

METADATOS ZOTERO:
{metadatos_texto}

TEXTOS FUENTE:
{fuentes}
"""


def prompt_referencias(fuentes, metadatos_texto, estilo_citacion):
    return f"""
Organiza las referencias bibliográficas reales en estilo {estilo_citacion}.

REGLAS:
- Usa prioritariamente los metadatos de Zotero.
- Complementa con los textos fuente si es necesario.
- No inventes DOI, revistas, autores, años ni editoriales.
- Si falta información, marca: [Referencia incompleta].
- Presenta solo referencias identificables.

METADATOS ZOTERO:
{metadatos_texto}

TEXTOS FUENTE:
{fuentes}
"""


def prompt_articulo_completo(
    titulo,
    tema,
    objetivo,
    metodologia,
    plantilla,
    estilo_citacion,
    fuentes,
    matriz,
    referencias,
    metadatos_texto,
    instrucciones,
    extension
):
    return f"""
Actúa como asistente académico experto en redacción científica.

OBJETIVO:
Redactar un borrador de artículo científico en español, basado únicamente en
las fuentes cargadas, metadatos de Zotero y matriz bibliográfica.

REGLAS OBLIGATORIAS:
- No inventes referencias, autores, años, revistas, DOI, cifras ni resultados.
- Usa información respaldada en fuentes.
- Parafrasea, no copies párrafos extensos.
- Usa citas en el texto según {estilo_citacion} cuando sea posible.
- Si falta información, escribe: [Información insuficiente en las fuentes].
- Mantén coherencia entre problema, objetivo, metodología, resultados y conclusiones.
- El resultado es un borrador académico revisable.

DATOS DEL ARTÍCULO:
Título tentativo: {titulo}
Tema central: {tema}
Objetivo: {objetivo}
Metodología esperada: {metodologia}
Extensión aproximada: {extension} palabras por sección.
Estilo de citación: {estilo_citacion}

PLANTILLA:
{plantilla}

INSTRUCCIONES ADICIONALES:
{instrucciones}

METADATOS ZOTERO:
{metadatos_texto}

MATRIZ BIBLIOGRÁFICA:
{matriz}

REFERENCIAS ORGANIZADAS:
{referencias}

TEXTOS FUENTE:
{fuentes}

TAREA:
Redacta el artículo completo siguiendo la plantilla.
Incluye título, resumen, palabras clave, introducción, revisión de literatura,
metodología, resultados o hallazgos, discusión, conclusiones, limitaciones,
futuras investigaciones y referencias.
"""


def prompt_generar_seccion(
    seccion,
    titulo,
    tema,
    objetivo,
    metodologia,
    estilo_citacion,
    fuentes,
    matriz,
    referencias,
    metadatos_texto,
    instrucciones,
    extension
):
    return f"""
Redacta únicamente la sección: {seccion}

Título: {titulo}
Tema: {tema}
Objetivo: {objetivo}
Metodología esperada: {metodologia}
Estilo de citación: {estilo_citacion}
Extensión aproximada: {extension} palabras.

REGLAS:
- Usa únicamente información respaldada en las fuentes.
- No inventes datos ni referencias.
- Usa tono académico.
- Cita autores cuando sea posible.
- No redactes otras secciones.
- Si no hay información suficiente, dilo explícitamente.

METADATOS ZOTERO:
{metadatos_texto}

MATRIZ BIBLIOGRÁFICA:
{matriz}

REFERENCIAS:
{referencias}

INSTRUCCIONES ADICIONALES:
{instrucciones}

TEXTOS FUENTE:
{fuentes}
"""


def prompt_evaluacion_calidad(articulo):
    return f"""
Evalúa la calidad académica del siguiente borrador de artículo científico.

Entrega una rúbrica con puntuación de 1 a 5 para:

1. Originalidad del tema
2. Claridad del problema
3. Coherencia del objetivo
4. Solidez metodológica
5. Uso de literatura científica
6. Calidad argumentativa
7. Coherencia entre resultados y conclusiones
8. Calidad de redacción académica
9. Calidad de citas y referencias
10. Potencial para envío a revista académica

Luego entrega:
- Fortalezas
- Debilidades
- Correcciones prioritarias
- Recomendaciones de mejora
- Nivel actual: Borrador inicial / Borrador avanzado / Candidato a revisión editorial

ARTÍCULO:
{articulo}
"""


def prompt_revision_coherencia(articulo):
    return f"""
Realiza una revisión de coherencia interna del artículo.

Compara:
1. Título vs contenido
2. Resumen vs desarrollo
3. Problema vs objetivo
4. Objetivo vs metodología
5. Metodología vs resultados
6. Resultados vs discusión
7. Discusión vs conclusiones
8. Citas vs referencias
9. Secciones faltantes

Entrega una tabla:
Elemento evaluado | Diagnóstico | Problema encontrado | Recomendación

ARTÍCULO:
{articulo}
"""


def prompt_resumen_abstract(articulo):
    return f"""
A partir del artículo, genera:

1. Resumen estructurado en español:
   - Objetivo
   - Metodología
   - Resultados
   - Conclusiones

2. Palabras clave en español, entre 4 y 6.

3. Title en inglés.

4. Abstract académico en inglés:
   - Objective
   - Methodology
   - Results
   - Conclusions

5. Keywords en inglés.

No inventes resultados.

ARTÍCULO:
{articulo}
"""


def prompt_parafrasis(texto):
    return f"""
Reescribe el texto con paráfrasis académica.

REGLAS:
- Conserva el sentido original.
- Mejora claridad, cohesión y formalidad.
- No inventes datos.
- No agregues referencias nuevas.
- Evita copia literal extensa.

TEXTO:
{texto}
"""


# =====================================================
# PROYECTOS
# =====================================================

def guardar_proyecto(nombre):
    carpeta = "proyectos_guardados"
    os.makedirs(carpeta, exist_ok=True)

    nombre_limpio = re.sub(r"[^a-zA-Z0-9_-]", "_", nombre)
    ruta = os.path.join(carpeta, f"{nombre_limpio}.json")

    datos = {
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "fuentes_extraidas": st.session_state.fuentes_extraidas,
        "metadatos_zotero": st.session_state.metadatos_zotero,
        "matriz_bibliografica": st.session_state.matriz_bibliografica,
        "referencias_extraidas": st.session_state.referencias_extraidas,
        "articulo_completo": st.session_state.articulo_completo,
        "evaluacion_calidad": st.session_state.evaluacion_calidad,
        "revision_coherencia": st.session_state.revision_coherencia,
        "resumen_abstract": st.session_state.resumen_abstract,
        "secciones_generadas": st.session_state.secciones_generadas,
        "datos_articulo": st.session_state.datos_articulo,
        "map_keywords": st.session_state.map_keywords,
        "network_keywords": st.session_state.network_keywords,
        "map_authors": st.session_state.map_authors,
        "network_authors": st.session_state.network_authors
    }

    with open(ruta, "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=4)

    return ruta


def cargar_proyecto(ruta):
    with open(ruta, "r", encoding="utf-8") as f:
        datos = json.load(f)

    for clave, valor in datos.items():
        st.session_state[clave] = valor


def listar_proyectos():
    carpeta = "proyectos_guardados"
    os.makedirs(carpeta, exist_ok=True)

    return [
        archivo for archivo in os.listdir(carpeta)
        if archivo.endswith(".json")
    ]


# =====================================================
# SIDEBAR
# =====================================================

with st.sidebar:
    st.header("⚙️ Configuración")

    modelo = st.selectbox(
        "Modelo OpenAI",
        ["gpt-4o-mini", "gpt-4o"],
        index=0
    )

    estilo_citacion = st.selectbox(
        "Estilo de citación",
        ["APA 7", "IEEE", "Vancouver", "ICONTEC"],
        index=0
    )

    extension_seccion = st.slider(
        "Palabras aproximadas por sección",
        min_value=150,
        max_value=1500,
        value=500,
        step=50
    )

    max_caracteres = st.slider(
        "Máximo de caracteres de fuentes",
        min_value=20000,
        max_value=180000,
        value=90000,
        step=10000
    )

    min_ocurrencias = st.slider(
        "Mínimo de ocurrencias para redes VOSviewer",
        min_value=1,
        max_value=10,
        value=1,
        step=1
    )


# =====================================================
# TABS
# =====================================================

tabs = st.tabs([
    "1. PDF manual",
    "2. Zotero",
    "3. Bibliometría VOSviewer",
    "4. Matriz y referencias",
    "5. Plantilla y datos",
    "6. Generar artículo",
    "7. Revisión académica",
    "8. Resumen / Abstract",
    "9. Exportar",
    "10. Proyectos"
])


# =====================================================
# TAB 1: PDF MANUAL
# =====================================================

with tabs[0]:
    st.header("1. Cargar PDF manualmente")

    archivos_pdf = st.file_uploader(
        "Suba uno o varios artículos científicos en PDF",
        type=["pdf"],
        accept_multiple_files=True
    )

    if archivos_pdf:
        st.success(f"Se cargaron {len(archivos_pdf)} archivo(s).")

        for archivo in archivos_pdf:
            st.write(f"- {archivo.name}")

        if st.button("Extraer texto de PDF manuales"):
            texto_total = ""

            with st.spinner("Extrayendo texto..."):
                for archivo in archivos_pdf:
                    texto_pdf = extraer_texto_pdf(archivo)

                    texto_total += "\n\n==============================\n"
                    texto_total += f"FUENTE PDF MANUAL: {archivo.name}\n"
                    texto_total += "==============================\n"
                    texto_total += texto_pdf

            st.session_state.fuentes_extraidas += "\n\n" + recortar_texto(
                texto_total,
                max_caracteres
            )

            st.success("Texto extraído y agregado a las fuentes.")

    if st.session_state.fuentes_extraidas:
        st.subheader("Fuentes extraídas acumuladas")

        st.text_area(
            "Texto fuente",
            value=st.session_state.fuentes_extraidas,
            height=500
        )


# =====================================================
# TAB 2: ZOTERO
# =====================================================

with tabs[1]:
    st.header("2. Conexión con Zotero")

    st.markdown(
        """
        Para usar Zotero necesita:
        - Zotero User ID o Group ID.
        - API Key de Zotero.
        - Permisos de lectura sobre la biblioteca.
        - Permiso de acceso a archivos si desea descargar PDF adjuntos.
        """
    )

    col1, col2 = st.columns(2)

    with col1:
        zotero_library_type = st.selectbox(
            "Tipo de biblioteca Zotero",
            ["user", "group"],
            index=0
        )

        zotero_library_id = st.text_input(
            "User ID o Group ID de Zotero"
        )

    with col2:
        zotero_api_key = st.text_input(
            "API Key de Zotero",
            type="password"
        )

        descargar_pdfs_zotero = st.checkbox(
            "Intentar descargar PDF adjuntos desde Zotero",
            value=True
        )

    if st.button("Listar colecciones de Zotero"):
        if not zotero_library_id or not zotero_api_key:
            st.warning("Ingrese Library ID y API Key de Zotero.")
        else:
            try:
                colecciones = zotero_listar_colecciones(
                    zotero_library_type,
                    zotero_library_id,
                    zotero_api_key
                )

                st.session_state.zotero_colecciones = colecciones

                st.success(f"Se encontraron {len(colecciones)} colecciones.")

            except Exception as e:
                st.error(f"No se pudo conectar con Zotero: {e}")

    colecciones = st.session_state.get("zotero_colecciones", [])

    if colecciones:
        opciones = {
            f"{c['name']} ({c['numItems']} ítems)": c["key"]
            for c in colecciones
        }

        seleccion = st.selectbox(
            "Seleccione una colección de Zotero",
            list(opciones.keys())
        )

        collection_key = opciones[seleccion]

        if st.button("Importar colección Zotero"):
            try:
                with st.spinner("Importando ítems desde Zotero..."):
                    items = zotero_items_coleccion(
                        zotero_library_type,
                        zotero_library_id,
                        zotero_api_key,
                        collection_key
                    )

                    metadatos = []
                    texto_total = ""

                    for item in items:
                        metadato = item_zotero_a_metadato(item)
                        metadatos.append(metadato)

                        if descargar_pdfs_zotero:
                            item_key = metadato.get("key", "")

                            try:
                                hijos = zotero_children(
                                    zotero_library_type,
                                    zotero_library_id,
                                    zotero_api_key,
                                    item_key
                                )

                                for hijo in hijos:
                                    data_hijo = hijo.get("data", {})
                                    content_type = data_hijo.get("contentType", "")
                                    title_hijo = data_hijo.get("title", "")
                                    key_hijo = data_hijo.get("key", hijo.get("key", ""))

                                    es_pdf = (
                                        "pdf" in content_type.lower()
                                        or title_hijo.lower().endswith(".pdf")
                                    )

                                    if es_pdf:
                                        archivo_pdf = zotero_descargar_pdf_adjunto(
                                            zotero_library_type,
                                            zotero_library_id,
                                            zotero_api_key,
                                            key_hijo
                                        )

                                        if archivo_pdf:
                                            texto_pdf = extraer_texto_pdf(archivo_pdf)

                                            texto_total += "\n\n==============================\n"
                                            texto_total += f"PDF ZOTERO: {metadato.get('title', '')}\n"
                                            texto_total += "==============================\n"
                                            texto_total += texto_pdf

                            except Exception:
                                pass

                    st.session_state.metadatos_zotero = metadatos

                    metadatos_texto = metadatos_a_texto(metadatos)

                    st.session_state.fuentes_extraidas += "\n\n"
                    st.session_state.fuentes_extraidas += metadatos_texto

                    if texto_total:
                        st.session_state.fuentes_extraidas += "\n\n"
                        st.session_state.fuentes_extraidas += recortar_texto(
                            texto_total,
                            max_caracteres
                        )

                st.success(
                    f"Se importaron {len(st.session_state.metadatos_zotero)} registros Zotero."
                )

            except Exception as e:
                st.error(f"Error importando colección Zotero: {e}")

    if st.session_state.metadatos_zotero:
        st.subheader("Metadatos importados desde Zotero")

        df_zotero = metadatos_a_dataframe(st.session_state.metadatos_zotero)

        st.dataframe(df_zotero, use_container_width=True)

        excel_zotero = BytesIO()
        with pd.ExcelWriter(excel_zotero, engine="openpyxl") as writer:
            df_zotero.to_excel(writer, index=False, sheet_name="Zotero")
        excel_zotero.seek(0)

        st.download_button(
            "📊 Descargar metadatos Zotero en Excel",
            data=excel_zotero,
            file_name="metadatos_zotero.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# =====================================================
# TAB 3: VOSVIEWER
# =====================================================

with tabs[2]:
    st.header("3. Bibliometría y redes para VOSviewer")

    st.markdown(
        """
        Esta sección genera archivos compatibles con VOSviewer:
        - `map_keywords.txt`
        - `network_keywords.txt`
        - `map_authors.txt`
        - `network_authors.txt`

        Los archivos se pueden abrir en VOSviewer como red bibliométrica.
        """
    )

    if not st.session_state.metadatos_zotero:
        st.warning("Primero importe una colección desde Zotero.")
    else:
        col1, col2 = st.columns(2)

        with col1:
            if st.button("Generar red de palabras clave"):
                map_txt, network_txt = generar_red_keywords_vosviewer(
                    st.session_state.metadatos_zotero,
                    min_ocurrencias=min_ocurrencias
                )

                st.session_state.map_keywords = map_txt
                st.session_state.network_keywords = network_txt

                st.success("Red de palabras clave generada.")

        with col2:
            if st.button("Generar red de coautoría"):
                map_txt, network_txt = generar_red_autores_vosviewer(
                    st.session_state.metadatos_zotero,
                    min_ocurrencias=min_ocurrencias
                )

                st.session_state.map_authors = map_txt
                st.session_state.network_authors = network_txt

                st.success("Red de coautoría generada.")

        st.subheader("Map keywords")

        st.text_area(
            "map_keywords.txt",
            value=st.session_state.map_keywords,
            height=250
        )

        st.subheader("Network keywords")

        st.text_area(
            "network_keywords.txt",
            value=st.session_state.network_keywords,
            height=250
        )

        st.subheader("Map authors")

        st.text_area(
            "map_authors.txt",
            value=st.session_state.map_authors,
            height=250
        )

        st.subheader("Network authors")

        st.text_area(
            "network_authors.txt",
            value=st.session_state.network_authors,
            height=250
        )


# =====================================================
# TAB 4: MATRIZ Y REFERENCIAS
# =====================================================

with tabs[3]:
    st.header("4. Matriz bibliográfica y referencias")

    metadatos_texto = metadatos_a_texto(st.session_state.metadatos_zotero)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Generar matriz bibliográfica"):
            if not st.session_state.fuentes_extraidas and not st.session_state.metadatos_zotero:
                st.warning("Primero cargue fuentes o importe Zotero.")
            else:
                with st.spinner("Generando matriz..."):
                    prompt = prompt_matriz_bibliografica(
                        st.session_state.fuentes_extraidas,
                        metadatos_texto,
                        estilo_citacion
                    )

                    st.session_state.matriz_bibliografica = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Matriz generada.")

    with col2:
        if st.button("Organizar referencias"):
            if not st.session_state.fuentes_extraidas and not st.session_state.metadatos_zotero:
                st.warning("Primero cargue fuentes o importe Zotero.")
            else:
                with st.spinner("Organizando referencias..."):
                    prompt = prompt_referencias(
                        st.session_state.fuentes_extraidas,
                        metadatos_texto,
                        estilo_citacion
                    )

                    st.session_state.referencias_extraidas = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.1
                    )

                st.success("Referencias organizadas.")

    st.subheader("Matriz bibliográfica")

    st.text_area(
        "Matriz",
        value=st.session_state.matriz_bibliografica,
        height=450
    )

    st.subheader("Referencias organizadas")

    st.text_area(
        "Referencias",
        value=st.session_state.referencias_extraidas,
        height=300
    )


# =====================================================
# TAB 5: PLANTILLA Y DATOS
# =====================================================

with tabs[4]:
    st.header("5. Datos del artículo y plantilla")

    titulo = st.text_input(
        "Título tentativo",
        value=st.session_state.datos_articulo.get("titulo", "")
    )

    tema = st.text_area(
        "Tema central",
        value=st.session_state.datos_articulo.get("tema", "")
    )

    objetivo = st.text_area(
        "Objetivo del artículo",
        value=st.session_state.datos_articulo.get("objetivo", "")
    )

    metodologia = st.text_area(
        "Metodología esperada",
        value=st.session_state.datos_articulo.get("metodologia", "")
    )

    tipo_plantilla = st.selectbox(
        "Tipo de artículo",
        [
            "Artículo empírico tipo revista indexada",
            "Artículo de revisión",
            "Artículo IMRyD",
            "Artículo de reflexión",
            "Plantilla personalizada"
        ]
    )

    if tipo_plantilla == "Plantilla personalizada":
        plantilla = st.text_area(
            "Plantilla personalizada",
            value=st.session_state.datos_articulo.get(
                "plantilla",
                "1. Título\n2. Resumen\n3. Palabras clave\n4. Introducción\n5. Metodología\n6. Resultados\n7. Discusión\n8. Conclusiones\n9. Referencias"
            ),
            height=350
        )
    else:
        plantilla = obtener_plantilla(tipo_plantilla)

        st.text_area(
            "Plantilla seleccionada",
            value=plantilla,
            height=350
        )

    instrucciones_adicionales = st.text_area(
        "Instrucciones adicionales",
        value=st.session_state.datos_articulo.get("instrucciones_adicionales", "")
    )

    if st.button("Guardar datos del artículo"):
        st.session_state.datos_articulo = {
            "titulo": titulo,
            "tema": tema,
            "objetivo": objetivo,
            "metodologia": metodologia,
            "tipo_plantilla": tipo_plantilla,
            "plantilla": plantilla,
            "instrucciones_adicionales": instrucciones_adicionales
        }

        st.success("Datos del artículo guardados.")


# =====================================================
# TAB 6: GENERAR ARTICULO
# =====================================================

with tabs[5]:
    st.header("6. Generar artículo")

    datos = st.session_state.datos_articulo
    metadatos_texto = metadatos_a_texto(st.session_state.metadatos_zotero)

    modo = st.radio(
        "Modo de generación",
        ["Generar artículo completo", "Generar por secciones"],
        horizontal=True
    )

    secciones = [
        "Resumen",
        "Palabras clave",
        "Introducción",
        "Revisión de literatura",
        "Marco teórico",
        "Metodología",
        "Resultados",
        "Discusión",
        "Conclusiones",
        "Limitaciones y futuras investigaciones",
        "Referencias"
    ]

    if modo == "Generar artículo completo":
        if st.button("Generar artículo completo"):
            if not datos.get("titulo"):
                st.warning("Primero guarde los datos del artículo.")
            elif not st.session_state.fuentes_extraidas and not st.session_state.metadatos_zotero:
                st.warning("Primero cargue fuentes PDF o importe Zotero.")
            else:
                with st.spinner("Generando artículo completo..."):
                    prompt = prompt_articulo_completo(
                        titulo=datos["titulo"],
                        tema=datos["tema"],
                        objetivo=datos["objetivo"],
                        metodologia=datos["metodologia"],
                        plantilla=datos["plantilla"],
                        estilo_citacion=estilo_citacion,
                        fuentes=st.session_state.fuentes_extraidas,
                        matriz=st.session_state.matriz_bibliografica,
                        referencias=st.session_state.referencias_extraidas,
                        metadatos_texto=metadatos_texto,
                        instrucciones=datos["instrucciones_adicionales"],
                        extension=extension_seccion
                    )

                    st.session_state.articulo_completo = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.35
                    )

                st.success("Artículo generado.")

    else:
        seccion = st.selectbox("Seleccione la sección", secciones)

        if st.button(f"Generar sección: {seccion}"):
            if not datos.get("titulo"):
                st.warning("Primero guarde los datos del artículo.")
            else:
                with st.spinner(f"Generando {seccion}..."):
                    prompt = prompt_generar_seccion(
                        seccion=seccion,
                        titulo=datos["titulo"],
                        tema=datos["tema"],
                        objetivo=datos["objetivo"],
                        metodologia=datos["metodologia"],
                        estilo_citacion=estilo_citacion,
                        fuentes=st.session_state.fuentes_extraidas,
                        matriz=st.session_state.matriz_bibliografica,
                        referencias=st.session_state.referencias_extraidas,
                        metadatos_texto=metadatos_texto,
                        instrucciones=datos["instrucciones_adicionales"],
                        extension=extension_seccion
                    )

                    resultado = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.35
                    )

                    st.session_state.secciones_generadas[seccion] = resultado

                texto_unido = ""

                for nombre, contenido in st.session_state.secciones_generadas.items():
                    texto_unido += f"\n\n{nombre.upper()}\n\n{contenido}"

                st.session_state.articulo_completo = texto_unido

                st.success(f"Sección {seccion} generada.")

    st.subheader("Borrador editable")

    articulo_editado = st.text_area(
        "Artículo",
        value=st.session_state.articulo_completo,
        height=800
    )

    st.session_state.articulo_completo = articulo_editado


# =====================================================
# TAB 7: REVISION ACADEMICA
# =====================================================

with tabs[6]:
    st.header("7. Revisión académica")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Evaluar calidad académica"):
            if not st.session_state.articulo_completo:
                st.warning("Primero genere o pegue un artículo.")
            else:
                with st.spinner("Evaluando calidad..."):
                    st.session_state.evaluacion_calidad = llamar_openai(
                        prompt_evaluacion_calidad(st.session_state.articulo_completo),
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Evaluación generada.")

    with col2:
        if st.button("Revisar coherencia interna"):
            if not st.session_state.articulo_completo:
                st.warning("Primero genere o pegue un artículo.")
            else:
                with st.spinner("Revisando coherencia..."):
                    st.session_state.revision_coherencia = llamar_openai(
                        prompt_revision_coherencia(st.session_state.articulo_completo),
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Revisión generada.")

    st.subheader("Evaluación de calidad")

    st.text_area(
        "Rúbrica",
        value=st.session_state.evaluacion_calidad,
        height=400
    )

    st.subheader("Revisión de coherencia")

    st.text_area(
        "Coherencia",
        value=st.session_state.revision_coherencia,
        height=400
    )

    st.subheader("Paráfrasis académica")

    texto_parafrasear = st.text_area(
        "Fragmento para parafrasear",
        height=200
    )

    if st.button("Parafrasear"):
        if not texto_parafrasear.strip():
            st.warning("Pegue un texto.")
        else:
            with st.spinner("Parafraseando..."):
                st.session_state.parafrasis = llamar_openai(
                    prompt_parafrasis(texto_parafrasear),
                    modelo=modelo,
                    temperatura=0.25
                )

    st.text_area(
        "Resultado de paráfrasis",
        value=st.session_state.parafrasis,
        height=250
    )


# =====================================================
# TAB 8: RESUMEN / ABSTRACT
# =====================================================

with tabs[7]:
    st.header("8. Resumen estructurado, abstract y keywords")

    if st.button("Generar resumen y abstract"):
        if not st.session_state.articulo_completo:
            st.warning("Primero genere o pegue un artículo.")
        else:
            with st.spinner("Generando resumen y abstract..."):
                st.session_state.resumen_abstract = llamar_openai(
                    prompt_resumen_abstract(st.session_state.articulo_completo),
                    modelo=modelo,
                    temperatura=0.25
                )

            st.success("Resumen y abstract generados.")

    st.text_area(
        "Resumen / Abstract",
        value=st.session_state.resumen_abstract,
        height=600
    )


# =====================================================
# TAB 9: EXPORTAR
# =====================================================

with tabs[8]:
    st.header("9. Exportar resultados")

    articulo = st.session_state.articulo_completo

    if articulo:
        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "📄 Descargar artículo PDF",
                data=generar_pdf(articulo),
                file_name="articulo_cientifico.pdf",
                mime="application/pdf"
            )

        with col2:
            st.download_button(
                "📝 Descargar artículo Word",
                data=generar_docx(articulo, "Artículo científico generado"),
                file_name="articulo_cientifico.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col3:
            st.download_button(
                "📃 Descargar artículo TXT",
                data=articulo.encode("utf-8"),
                file_name="articulo_cientifico.txt",
                mime="text/plain"
            )
    else:
        st.warning("Aún no hay artículo para exportar.")

    st.subheader("Exportar matriz bibliográfica")

    if st.session_state.matriz_bibliografica:
        st.download_button(
            "📊 Descargar matriz Excel",
            data=generar_excel_desde_texto(st.session_state.matriz_bibliografica),
            file_name="matriz_bibliografica.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.download_button(
            "📃 Descargar matriz TXT",
            data=st.session_state.matriz_bibliografica.encode("utf-8"),
            file_name="matriz_bibliografica.txt",
            mime="text/plain"
        )

    st.subheader("Exportar archivos VOSviewer")

    col1, col2 = st.columns(2)

    with col1:
        if st.session_state.map_keywords:
            st.download_button(
                "Descargar map_keywords.txt",
                data=st.session_state.map_keywords.encode("utf-8"),
                file_name="map_keywords.txt",
                mime="text/plain"
            )

        if st.session_state.network_keywords:
            st.download_button(
                "Descargar network_keywords.txt",
                data=st.session_state.network_keywords.encode("utf-8"),
                file_name="network_keywords.txt",
                mime="text/plain"
            )

    with col2:
        if st.session_state.map_authors:
            st.download_button(
                "Descargar map_authors.txt",
                data=st.session_state.map_authors.encode("utf-8"),
                file_name="map_authors.txt",
                mime="text/plain"
            )

        if st.session_state.network_authors:
            st.download_button(
                "Descargar network_authors.txt",
                data=st.session_state.network_authors.encode("utf-8"),
                file_name="network_authors.txt",
                mime="text/plain"
            )


# =====================================================
# TAB 10: PROYECTOS
# =====================================================

with tabs[9]:
    st.header("10. Guardado y carga de proyectos")

    nombre_proyecto = st.text_input(
        "Nombre del proyecto",
        value=st.session_state.proyecto_actual
    )

    if st.button("Guardar proyecto"):
        if not nombre_proyecto.strip():
            st.warning("Ingrese un nombre para el proyecto.")
        else:
            ruta = guardar_proyecto(nombre_proyecto)
            st.session_state.proyecto_actual = nombre_proyecto
            st.success(f"Proyecto guardado en: {ruta}")

    proyectos = listar_proyectos()

    if proyectos:
        proyecto = st.selectbox("Proyecto guardado", proyectos)

        if st.button("Cargar proyecto"):
            ruta = os.path.join("proyectos_guardados", proyecto)
            cargar_proyecto(ruta)
            st.success("Proyecto cargado correctamente.")
    else:
        st.info("Aún no hay proyectos guardados.")

    st.markdown(
        """
        **Flujo profesional recomendado:**

        1. Conectar Zotero o cargar PDF manualmente.
        2. Importar metadatos y PDF.
        3. Generar red de palabras clave y coautoría para VOSviewer.
        4. Generar matriz bibliográfica.
        5. Organizar referencias.
        6. Completar título, tema, objetivo y metodología.
        7. Generar el artículo por secciones.
        8. Revisar calidad académica.
        9. Revisar coherencia interna.
        10. Generar resumen, abstract y keywords.
        11. Exportar Word, PDF, Excel y archivos VOSviewer.
        """
    )