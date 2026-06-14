import streamlit as st
from openai import OpenAI, RateLimitError, AuthenticationError
from pypdf import PdfReader
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import pandas as pd
import os
import json
import re
from datetime import datetime


# =====================================================
# CONFIGURACIÓN GENERAL
# =====================================================

st.set_page_config(
    page_title="Asistente Científico de Artículos",
    layout="wide"
)

st.title("📚 Asistente Científico para Elaboración de Artículos Académicos")

st.info(
    "Esta aplicación permite cargar artículos científicos en PDF, extraer información, "
    "generar una matriz bibliográfica, construir artículos por secciones, revisar calidad "
    "académica, verificar coherencia, generar resumen en español e inglés y exportar en "
    "Word, PDF, Excel y TXT."
)


# =====================================================
# CONFIGURACIÓN API
# =====================================================

try:
    api_key = st.secrets["OPENAI_API_KEY"]
except Exception:
    api_key = None

if not api_key:
    st.error(
        "No se encontró la clave API de OpenAI. "
        "Cree el archivo .streamlit/secrets.toml con OPENAI_API_KEY."
    )
    st.stop()

client = OpenAI(api_key=api_key)


# =====================================================
# ESTADO DE SESIÓN
# =====================================================

valores_iniciales = {
    "fuentes_extraidas": "",
    "matriz_bibliografica": "",
    "referencias_extraidas": "",
    "articulo_completo": "",
    "evaluacion_calidad": "",
    "revision_coherencia": "",
    "resumen_estructurado": "",
    "abstract_ingles": "",
    "parafrasis": "",
    "proyecto_actual": "",
}

for clave, valor in valores_iniciales.items():
    if clave not in st.session_state:
        st.session_state[clave] = valor

if "secciones_generadas" not in st.session_state:
    st.session_state.secciones_generadas = {}

if "historial_proyectos" not in st.session_state:
    st.session_state.historial_proyectos = []


# =====================================================
# FUNCIONES GENERALES
# =====================================================

def llamar_openai(prompt, modelo="gpt-4o-mini", temperatura=0.3):
    """
    Llama al modelo de OpenAI y maneja errores frecuentes.
    """
    try:
        respuesta = client.chat.completions.create(
            model=modelo,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un asistente académico experto en redacción científica, "
                        "metodología de investigación, revisión bibliográfica y normas de citación. "
                        "Debes evitar inventar datos, referencias, autores, DOI o resultados."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=temperatura
        )

        return respuesta.choices[0].message.content

    except AuthenticationError:
        return "ERROR: La clave API de OpenAI no es válida. Revise .streamlit/secrets.toml."

    except RateLimitError:
        return (
            "ERROR: La cuenta de OpenAI no tiene cuota disponible o superó el límite de uso. "
            "Revise Billing y Usage Limits en OpenAI Platform."
        )

    except Exception as e:
        return f"ERROR: Ocurrió un problema al llamar al modelo: {e}"


def extraer_texto_pdf(archivo_pdf):
    """
    Extrae texto de un PDF. Funciona mejor con PDF que tienen texto seleccionable.
    """
    texto = ""

    try:
        lector = PdfReader(archivo_pdf)

        for numero_pagina, pagina in enumerate(lector.pages, start=1):
            contenido = pagina.extract_text()

            if contenido:
                texto += f"\n\n--- Página {numero_pagina} ---\n"
                texto += contenido

    except Exception as e:
        texto = f"Error al leer PDF: {e}"

    return texto


def recortar_texto(texto, max_caracteres):
    """
    Recorta texto para evitar superar límites del modelo.
    """
    if len(texto) > max_caracteres:
        return (
            texto[:max_caracteres]
            + "\n\n[Texto recortado automáticamente por longitud. "
            + "Para mayor precisión, cargue menos PDF o aumente el límite si su modelo lo permite.]"
        )

    return texto


def limpiar_texto_pdf(texto):
    """
    Limpia caracteres que pueden generar problemas en PDF básico.
    """
    reemplazos = {
        "📚": "",
        "📄": "",
        "✅": "",
        "✓": "",
        "✔": "",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "–": "-",
        "—": "-",
        "•": "-",
        "…": "...",
        "α": "alpha",
        "β": "beta",
        "γ": "gamma",
    }

    for original, reemplazo in reemplazos.items():
        texto = texto.replace(original, reemplazo)

    return texto


def extraer_bloque_referencias(texto):
    """
    Intenta extraer la sección de referencias de los textos fuente.
    """
    patrones = [
        r"references\s*(.*)",
        r"referencias\s*(.*)",
        r"bibliography\s*(.*)",
        r"bibliografía\s*(.*)"
    ]

    referencias = []

    for patron in patrones:
        coincidencias = re.findall(patron, texto, flags=re.IGNORECASE | re.DOTALL)
        for c in coincidencias:
            referencias.append(c[:12000])

    if referencias:
        return "\n\n--- REFERENCIAS DETECTADAS ---\n\n".join(referencias)

    return "No se detectó automáticamente una sección clara de referencias en los PDF."


def generar_pdf(texto):
    """
    Genera PDF en memoria.
    """
    texto = limpiar_texto_pdf(texto)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)

    for linea in texto.split("\n"):
        if linea.strip() == "":
            pdf.ln(4)
        else:
            try:
                pdf.multi_cell(0, 8, linea.encode("latin-1", errors="ignore").decode("latin-1"))
            except Exception:
                pdf.multi_cell(0, 8, "[Texto omitido por caracteres incompatibles]")

    pdf_bytes = pdf.output(dest="S").encode("latin-1", errors="ignore")
    return pdf_bytes


def generar_docx(texto, titulo_documento="Articulo generado"):
    """
    Genera documento Word en memoria.
    """
    documento = Document()

    section = documento.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    estilos = documento.styles
    estilos["Normal"].font.name = "Times New Roman"
    estilos["Normal"].font.size = Pt(12)

    documento.add_heading(titulo_documento, level=0)

    for linea in texto.split("\n"):
        linea_limpia = linea.strip()

        if linea_limpia == "":
            documento.add_paragraph("")
        elif re.match(r"^\d+[\.\)]\s+", linea_limpia):
            documento.add_heading(linea_limpia, level=1)
        elif linea_limpia.lower() in [
            "resumen", "abstract", "introducción", "introduccion",
            "metodología", "metodologia", "resultados", "discusión",
            "discusion", "conclusiones", "referencias"
        ]:
            documento.add_heading(linea_limpia, level=1)
        else:
            documento.add_paragraph(linea_limpia)

    archivo = BytesIO()
    documento.save(archivo)
    archivo.seek(0)
    return archivo


def generar_excel_matriz(texto_matriz):
    """
    Genera Excel simple con la matriz bibliográfica en una sola hoja.
    """
    filas = []

    for linea in texto_matriz.split("\n"):
        if linea.strip():
            filas.append({"Contenido": linea.strip()})

    df = pd.DataFrame(filas)

    archivo = BytesIO()
    with pd.ExcelWriter(archivo, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Matriz bibliografica")

    archivo.seek(0)
    return archivo


def guardar_proyecto(nombre, datos):
    """
    Guarda el proyecto en archivo JSON local.
    """
    carpeta = "proyectos_guardados"
    os.makedirs(carpeta, exist_ok=True)

    nombre_limpio = re.sub(r"[^a-zA-Z0-9_-]", "_", nombre)
    ruta = os.path.join(carpeta, f"{nombre_limpio}.json")

    with open(ruta, "w", encoding="utf-8") as f:
        json.dump(datos, f, ensure_ascii=False, indent=4)

    return ruta


def cargar_proyecto(ruta):
    """
    Carga un proyecto JSON.
    """
    with open(ruta, "r", encoding="utf-8") as f:
        return json.load(f)


def listar_proyectos():
    carpeta = "proyectos_guardados"
    os.makedirs(carpeta, exist_ok=True)

    archivos = [
        archivo for archivo in os.listdir(carpeta)
        if archivo.endswith(".json")
    ]

    return archivos


# =====================================================
# PLANTILLAS
# =====================================================

def obtener_plantilla(nombre):
    plantillas = {
        "Artículo empírico tipo revista indexada": """
1. Título
2. Resumen
   - Problema
   - Objetivo
   - Metodología
   - Resultados principales
   - Conclusión central
3. Palabras clave
4. Introducción
   - Contextualización del problema
   - Justificación
   - Brecha investigativa
   - Objetivo
   - Organización del artículo
5. Revisión de literatura
   - Antecedentes
   - Enfoques teóricos
   - Estudios relacionados
   - Vacíos identificados
6. Metodología
   - Enfoque
   - Tipo de estudio
   - Fuentes o datos utilizados
   - Procedimiento
   - Técnicas de análisis
7. Resultados
8. Discusión
9. Conclusiones
10. Limitaciones y futuras investigaciones
11. Referencias
""",

        "Artículo de revisión": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Metodología de revisión
   - Criterios de inclusión
   - Criterios de exclusión
   - Fuentes consultadas
   - Procedimiento de análisis
6. Desarrollo temático
7. Tendencias de la literatura
8. Discusión
9. Conclusiones
10. Referencias
""",

        "Artículo IMRyD": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Metodología
6. Resultados
7. Discusión
8. Conclusiones
9. Referencias
""",

        "Artículo de reflexión": """
1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Planteamiento del problema
6. Desarrollo argumentativo
7. Discusión crítica
8. Conclusiones
9. Referencias
""",

        "Plantilla personalizada": ""
    }

    return plantillas.get(nombre, plantillas["Artículo empírico tipo revista indexada"])


# =====================================================
# PROMPTS
# =====================================================

def prompt_matriz_bibliografica(fuentes, estilo_citacion):
    return f"""
Analiza los textos fuente extraídos de artículos científicos en PDF.

Elabora una matriz bibliográfica rigurosa en español con esta estructura:

| Autor(es) | Año | Título | Objetivo | Metodología | Datos/Muestra | Resultados principales | Limitaciones | Aporte al tema | Referencia en {estilo_citacion} |

REGLAS:
- No inventes autores, años, revistas, DOI ni datos.
- Si un dato no aparece, escribe: No identificado en el PDF.
- Si una referencia está incompleta, escribe: Referencia incompleta.
- Usa lenguaje académico.
- Prioriza precisión sobre extensión.

TEXTOS FUENTE:
{fuentes}
"""


def prompt_referencias(fuentes, estilo_citacion):
    return f"""
Extrae y organiza las referencias bibliográficas identificables en los PDF.

Estilo solicitado: {estilo_citacion}

REGLAS:
- No inventes referencias.
- No inventes DOI.
- No inventes autores ni años.
- Si una referencia está incompleta, márcala como: [Referencia incompleta].
- Presenta solo referencias que aparezcan o puedan identificarse claramente en los textos fuente.

TEXTOS FUENTE:
{fuentes}
"""


def prompt_articulo_completo(
    titulo,
    tema,
    objetivo,
    metodologia,
    plantilla,
    estilo_citacion,
    fuentes,
    matriz,
    referencias,
    instrucciones,
    extension
):
    return f"""
Actúa como asistente académico experto en redacción de artículos científicos.

OBJETIVO:
Redactar un borrador de artículo científico en español, con alto nivel académico,
basado únicamente en los PDF cargados por el usuario.

REGLAS OBLIGATORIAS:
- Usa únicamente información respaldada en los textos fuente.
- No inventes referencias, autores, años, revistas, DOI, cifras ni resultados.
- No copies párrafos extensos de las fuentes; parafrasea.
- Usa citas en el texto según {estilo_citacion} cuando sea posible.
- Si falta información, escribe: [Información insuficiente en las fuentes cargadas].
- Si una referencia no se puede verificar, escribe: [Referencia no identificada claramente en los PDF].
- Mantén coherencia entre título, problema, objetivo, metodología, resultados y conclusiones.
- El texto debe ser un borrador académico revisable, no una promesa de publicación.

DATOS DEL ARTÍCULO:
Título tentativo: {titulo}
Tema central: {tema}
Objetivo: {objetivo}
Metodología esperada: {metodologia}
Extensión aproximada: {extension} palabras por sección.
Estilo de citación: {estilo_citacion}

PLANTILLA O ESTRUCTURA:
{plantilla}

INSTRUCCIONES ADICIONALES:
{instrucciones}

MATRIZ BIBLIOGRÁFICA PREVIA:
{matriz}

REFERENCIAS EXTRAÍDAS:
{referencias}

TEXTOS FUENTE:
{fuentes}

TAREA:
Redacta el artículo completo siguiendo la plantilla. Incluye:
- Título
- Resumen
- Palabras clave
- Introducción
- Revisión de literatura o marco teórico
- Metodología
- Resultados o hallazgos derivados de la revisión
- Discusión
- Conclusiones
- Limitaciones y futuras investigaciones
- Referencias
"""


def prompt_generar_seccion(
    nombre_seccion,
    titulo,
    tema,
    objetivo,
    metodologia,
    estilo_citacion,
    fuentes,
    matriz,
    referencias,
    instrucciones,
    extension
):
    return f"""
Redacta únicamente la sección: {nombre_seccion}

Artículo:
Título: {titulo}
Tema: {tema}
Objetivo: {objetivo}
Metodología esperada: {metodologia}
Estilo de citación: {estilo_citacion}
Extensión aproximada: {extension} palabras.

REGLAS:
- Usa únicamente información de los PDF.
- No inventes datos ni referencias.
- Usa tono académico.
- Cita autores cuando sea posible.
- No redactes otras secciones.
- Si no hay información suficiente, dilo explícitamente.

MATRIZ BIBLIOGRÁFICA:
{matriz}

REFERENCIAS EXTRAÍDAS:
{referencias}

INSTRUCCIONES ADICIONALES:
{instrucciones}

TEXTOS FUENTE:
{fuentes}
"""


def prompt_evaluacion_calidad(articulo):
    return f"""
Evalúa la calidad académica del siguiente borrador de artículo científico.

Entrega una rúbrica con puntuación de 1 a 5 para:

1. Originalidad del tema
2. Claridad del problema
3. Coherencia del objetivo
4. Solidez metodológica
5. Uso de literatura científica
6. Calidad de argumentación
7. Coherencia entre resultados y conclusiones
8. Calidad de redacción académica
9. Calidad de citas y referencias
10. Potencial para envío a revista académica

Luego entrega:
- Fortalezas
- Debilidades
- Correcciones prioritarias
- Recomendaciones para mejorar calidad científica
- Nivel actual: Borrador inicial / Borrador avanzado / Candidato a revisión editorial

ARTÍCULO:
{articulo}
"""


def prompt_revision_coherencia(articulo):
    return f"""
Realiza una revisión de coherencia interna del siguiente artículo.

Compara:

1. Título vs contenido
2. Resumen vs desarrollo
3. Problema vs objetivo
4. Objetivo vs metodología
5. Metodología vs resultados
6. Resultados vs discusión
7. Discusión vs conclusiones
8. Citas vs referencias
9. Secciones faltantes
10. Recomendaciones concretas de mejora

Usa una tabla con columnas:
Elemento evaluado | Diagnóstico | Problema encontrado | Recomendación

ARTÍCULO:
{articulo}
"""


def prompt_resumen_estructurado(articulo):
    return f"""
A partir del siguiente artículo, genera:

1. Resumen estructurado en español con:
   - Objetivo
   - Metodología
   - Resultados
   - Conclusiones

2. Palabras clave en español, entre 4 y 6.

3. Title en inglés.

4. Abstract académico en inglés con:
   - Objective
   - Methodology
   - Results
   - Conclusions

5. Keywords en inglés.

No inventes resultados. Usa solo lo que aparece en el artículo.

ARTÍCULO:
{articulo}
"""


def prompt_parafrasis(texto):
    return f"""
Reescribe el siguiente texto con paráfrasis académica.

REGLAS:
- Conserva el sentido original.
- Mejora la claridad, cohesión y formalidad.
- No inventes datos.
- No agregues referencias nuevas.
- Evita copiar literalmente frases extensas.

TEXTO:
{texto}
"""


# =====================================================
# SIDEBAR
# =====================================================

with st.sidebar:
    st.header("⚙️ Configuración")

    modelo = st.selectbox(
        "Modelo",
        ["gpt-4o-mini", "gpt-4o"],
        index=0
    )

    estilo_citacion = st.selectbox(
        "Estilo de citación",
        ["APA 7", "IEEE", "Vancouver", "ICONTEC"],
        index=0
    )

    extension_seccion = st.slider(
        "Palabras aproximadas por sección",
        min_value=150,
        max_value=1500,
        value=500,
        step=50
    )

    max_caracteres = st.slider(
        "Máximo de caracteres de fuentes",
        min_value=20000,
        max_value=180000,
        value=90000,
        step=10000
    )

    st.caption(
        "Si carga muchos PDF, se recomienda generar primero la matriz bibliográfica "
        "y luego construir el artículo por secciones."
    )


# =====================================================
# PESTAÑAS PRINCIPALES
# =====================================================

tabs = st.tabs([
    "1. Fuentes PDF",
    "2. Matriz bibliográfica",
    "3. Plantilla y datos",
    "4. Generar artículo",
    "5. Revisión académica",
    "6. Resumen / Abstract",
    "7. Exportar",
    "8. Proyectos"
])


# =====================================================
# TAB 1: FUENTES PDF
# =====================================================

with tabs[0]:
    st.header("1. Cargar y procesar fuentes PDF")

    archivos_pdf = st.file_uploader(
        "Suba uno o varios artículos científicos en PDF",
        type=["pdf"],
        accept_multiple_files=True
    )

    if archivos_pdf:
        st.success(f"Se cargaron {len(archivos_pdf)} archivo(s).")

        for archivo in archivos_pdf:
            st.write(f"- {archivo.name}")

        if st.button("Extraer texto de los PDF"):
            texto_total = ""

            with st.spinner("Extrayendo texto de los PDF..."):
                for archivo in archivos_pdf:
                    texto_pdf = extraer_texto_pdf(archivo)

                    texto_total += "\n\n==============================\n"
                    texto_total += f"FUENTE PDF: {archivo.name}\n"
                    texto_total += "==============================\n"
                    texto_total += texto_pdf

            st.session_state.fuentes_extraidas = recortar_texto(
                texto_total,
                max_caracteres
            )

            st.session_state.referencias_extraidas = extraer_bloque_referencias(
                st.session_state.fuentes_extraidas
            )

            st.success("Texto extraído correctamente.")

    if st.session_state.fuentes_extraidas:
        st.subheader("Texto extraído")

        st.text_area(
            "Contenido de las fuentes",
            value=st.session_state.fuentes_extraidas,
            height=400
        )

        st.subheader("Referencias detectadas automáticamente")

        st.text_area(
            "Bloque de referencias",
            value=st.session_state.referencias_extraidas,
            height=250
        )


# =====================================================
# TAB 2: MATRIZ BIBLIOGRÁFICA
# =====================================================

with tabs[1]:
    st.header("2. Matriz bibliográfica y referencias reales")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Generar matriz bibliográfica"):
            if not st.session_state.fuentes_extraidas:
                st.warning("Primero cargue y extraiga texto de los PDF.")
            else:
                with st.spinner("Generando matriz bibliográfica..."):
                    prompt = prompt_matriz_bibliografica(
                        st.session_state.fuentes_extraidas,
                        estilo_citacion
                    )

                    st.session_state.matriz_bibliografica = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Matriz generada.")

    with col2:
        if st.button("Organizar referencias"):
            if not st.session_state.fuentes_extraidas:
                st.warning("Primero cargue y extraiga texto de los PDF.")
            else:
                with st.spinner("Organizando referencias..."):
                    prompt = prompt_referencias(
                        st.session_state.fuentes_extraidas,
                        estilo_citacion
                    )

                    st.session_state.referencias_extraidas = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.1
                    )

                st.success("Referencias organizadas.")

    st.subheader("Matriz bibliográfica")

    st.text_area(
        "Matriz",
        value=st.session_state.matriz_bibliografica,
        height=450
    )

    st.subheader("Referencias organizadas")

    st.text_area(
        "Referencias",
        value=st.session_state.referencias_extraidas,
        height=300
    )


# =====================================================
# TAB 3: PLANTILLA Y DATOS
# =====================================================

with tabs[2]:
    st.header("3. Datos del artículo y plantilla")

    titulo = st.text_input(
        "Título tentativo",
        placeholder="Ejemplo: Inteligencia artificial y predicción de la deserción estudiantil en educación superior"
    )

    tema = st.text_area(
        "Tema central",
        placeholder="Describa el tema principal del artículo."
    )

    objetivo = st.text_area(
        "Objetivo del artículo",
        placeholder="Ejemplo: Analizar los aportes de los modelos de machine learning..."
    )

    metodologia = st.text_area(
        "Metodología esperada",
        placeholder="Ejemplo: Revisión documental, análisis comparativo, estudio empírico..."
    )

    tipo_plantilla = st.selectbox(
        "Tipo de artículo",
        [
            "Artículo empírico tipo revista indexada",
            "Artículo de revisión",
            "Artículo IMRyD",
            "Artículo de reflexión",
            "Plantilla personalizada"
        ]
    )

    if tipo_plantilla == "Plantilla personalizada":
        plantilla = st.text_area(
            "Pegue aquí su plantilla personalizada",
            value="""1. Título
2. Resumen
3. Palabras clave
4. Introducción
5. Revisión de literatura
6. Metodología
7. Resultados
8. Discusión
9. Conclusiones
10. Referencias""",
            height=350
        )
    else:
        plantilla = obtener_plantilla(tipo_plantilla)
        st.text_area(
            "Plantilla seleccionada",
            value=plantilla,
            height=350
        )

    instrucciones_adicionales = st.text_area(
        "Instrucciones adicionales",
        placeholder="Ejemplo: Enfatizar educación superior, contexto latinoamericano, implicaciones institucionales, etc."
    )

    st.session_state.datos_articulo = {
        "titulo": titulo,
        "tema": tema,
        "objetivo": objetivo,
        "metodologia": metodologia,
        "plantilla": plantilla,
        "instrucciones_adicionales": instrucciones_adicionales,
        "tipo_plantilla": tipo_plantilla,
        "estilo_citacion": estilo_citacion,
        "extension_seccion": extension_seccion
    }


# =====================================================
# TAB 4: GENERAR ARTÍCULO
# =====================================================

with tabs[3]:
    st.header("4. Generar artículo")

    datos = st.session_state.get("datos_articulo", {})

    if not datos:
        st.warning("Primero complete la pestaña 'Plantilla y datos'.")

    modo_generacion = st.radio(
        "Modo de generación",
        [
            "Generar artículo completo",
            "Generar por secciones"
        ],
        horizontal=True
    )

    secciones_disponibles = [
        "Resumen",
        "Palabras clave",
        "Introducción",
        "Revisión de literatura",
        "Marco teórico",
        "Metodología",
        "Resultados",
        "Discusión",
        "Conclusiones",
        "Limitaciones y futuras investigaciones",
        "Referencias"
    ]

    if modo_generacion == "Generar artículo completo":
        if st.button("Generar artículo completo"):
            if not st.session_state.fuentes_extraidas:
                st.warning("Primero cargue fuentes PDF.")
            elif not datos.get("titulo"):
                st.warning("Ingrese un título tentativo.")
            elif not datos.get("tema"):
                st.warning("Ingrese el tema central.")
            elif not datos.get("objetivo"):
                st.warning("Ingrese el objetivo.")
            else:
                with st.spinner("Generando artículo completo..."):
                    prompt = prompt_articulo_completo(
                        titulo=datos["titulo"],
                        tema=datos["tema"],
                        objetivo=datos["objetivo"],
                        metodologia=datos["metodologia"],
                        plantilla=datos["plantilla"],
                        estilo_citacion=estilo_citacion,
                        fuentes=st.session_state.fuentes_extraidas,
                        matriz=st.session_state.matriz_bibliografica,
                        referencias=st.session_state.referencias_extraidas,
                        instrucciones=datos["instrucciones_adicionales"],
                        extension=extension_seccion
                    )

                    st.session_state.articulo_completo = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.35
                    )

                st.success("Artículo generado.")

    else:
        seccion_elegida = st.selectbox(
            "Seleccione la sección a generar",
            secciones_disponibles
        )

        if st.button(f"Generar sección: {seccion_elegida}"):
            if not st.session_state.fuentes_extraidas:
                st.warning("Primero cargue fuentes PDF.")
            elif not datos.get("titulo"):
                st.warning("Ingrese un título tentativo.")
            else:
                with st.spinner(f"Generando {seccion_elegida}..."):
                    prompt = prompt_generar_seccion(
                        nombre_seccion=seccion_elegida,
                        titulo=datos["titulo"],
                        tema=datos["tema"],
                        objetivo=datos["objetivo"],
                        metodologia=datos["metodologia"],
                        estilo_citacion=estilo_citacion,
                        fuentes=st.session_state.fuentes_extraidas,
                        matriz=st.session_state.matriz_bibliografica,
                        referencias=st.session_state.referencias_extraidas,
                        instrucciones=datos["instrucciones_adicionales"],
                        extension=extension_seccion
                    )

                    resultado = llamar_openai(
                        prompt,
                        modelo=modelo,
                        temperatura=0.35
                    )

                    st.session_state.secciones_generadas[seccion_elegida] = resultado

                st.success(f"Sección {seccion_elegida} generada.")

        if st.session_state.secciones_generadas:
            texto_unido = ""

            for seccion, contenido in st.session_state.secciones_generadas.items():
                texto_unido += f"\n\n{seccion.upper()}\n\n{contenido}"

            st.session_state.articulo_completo = texto_unido

    st.subheader("Borrador del artículo")

    articulo_editado = st.text_area(
        "Puede editar manualmente el artículo aquí",
        value=st.session_state.articulo_completo,
        height=800
    )

    st.session_state.articulo_completo = articulo_editado


# =====================================================
# TAB 5: REVISIÓN ACADÉMICA
# =====================================================

with tabs[4]:
    st.header("5. Revisión académica y científica")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Evaluar calidad académica"):
            if not st.session_state.articulo_completo:
                st.warning("Primero genere o pegue un artículo.")
            else:
                with st.spinner("Evaluando calidad académica..."):
                    st.session_state.evaluacion_calidad = llamar_openai(
                        prompt_evaluacion_calidad(st.session_state.articulo_completo),
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Evaluación generada.")

    with col2:
        if st.button("Revisar coherencia interna"):
            if not st.session_state.articulo_completo:
                st.warning("Primero genere o pegue un artículo.")
            else:
                with st.spinner("Revisando coherencia..."):
                    st.session_state.revision_coherencia = llamar_openai(
                        prompt_revision_coherencia(st.session_state.articulo_completo),
                        modelo=modelo,
                        temperatura=0.2
                    )

                st.success("Revisión generada.")

    st.subheader("Rúbrica de calidad académica")

    st.text_area(
        "Evaluación",
        value=st.session_state.evaluacion_calidad,
        height=400
    )

    st.subheader("Revisión de coherencia")

    st.text_area(
        "Coherencia interna",
        value=st.session_state.revision_coherencia,
        height=400
    )

    st.subheader("Herramienta de paráfrasis académica")

    texto_parafrasear = st.text_area(
        "Pegue aquí un fragmento para mejorar o parafrasear",
        height=200
    )

    if st.button("Parafrasear fragmento"):
        if not texto_parafrasear.strip():
            st.warning("Pegue un texto para parafrasear.")
        else:
            with st.spinner("Parafraseando..."):
                st.session_state.parafrasis = llamar_openai(
                    prompt_parafrasis(texto_parafrasear),
                    modelo=modelo,
                    temperatura=0.25
                )

    st.text_area(
        "Resultado de paráfrasis",
        value=st.session_state.parafrasis,
        height=250
    )


# =====================================================
# TAB 6: RESUMEN / ABSTRACT
# =====================================================

with tabs[5]:
    st.header("6. Resumen estructurado, abstract y keywords")

    if st.button("Generar resumen, abstract y palabras clave"):
        if not st.session_state.articulo_completo:
            st.warning("Primero genere o pegue un artículo.")
        else:
            with st.spinner("Generando resumen y abstract..."):
                resultado = llamar_openai(
                    prompt_resumen_estructurado(st.session_state.articulo_completo),
                    modelo=modelo,
                    temperatura=0.25
                )

                st.session_state.resumen_estructurado = resultado

            st.success("Resumen y abstract generados.")

    st.text_area(
        "Resumen estructurado / Abstract",
        value=st.session_state.resumen_estructurado,
        height=600
    )


# =====================================================
# TAB 7: EXPORTAR
# =====================================================

with tabs[6]:
    st.header("7. Exportar resultados")

    articulo = st.session_state.articulo_completo

    if not articulo:
        st.warning("Aún no hay artículo para exportar.")
    else:
        nombre_base = "articulo_cientifico_generado"

        col1, col2, col3 = st.columns(3)

        with col1:
            pdf_data = generar_pdf(articulo)

            st.download_button(
                "📄 Descargar artículo en PDF",
                data=pdf_data,
                file_name=f"{nombre_base}.pdf",
                mime="application/pdf"
            )

        with col2:
            docx_data = generar_docx(articulo, "Artículo científico generado")

            st.download_button(
                "📝 Descargar artículo en Word",
                data=docx_data,
                file_name=f"{nombre_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col3:
            st.download_button(
                "📃 Descargar artículo en TXT",
                data=articulo.encode("utf-8"),
                file_name=f"{nombre_base}.txt",
                mime="text/plain"
            )

    st.subheader("Exportar matriz bibliográfica")

    if st.session_state.matriz_bibliografica:
        excel_data = generar_excel_matriz(st.session_state.matriz_bibliografica)

        st.download_button(
            "📊 Descargar matriz en Excel",
            data=excel_data,
            file_name="matriz_bibliografica.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        st.download_button(
            "📃 Descargar matriz en TXT",
            data=st.session_state.matriz_bibliografica.encode("utf-8"),
            file_name="matriz_bibliografica.txt",
            mime="text/plain"
        )
    else:
        st.info("Primero genere la matriz bibliográfica.")


# =====================================================
# TAB 8: PROYECTOS
# =====================================================

with tabs[7]:
    st.header("8. Gestión de proyectos")

    nombre_proyecto = st.text_input(
        "Nombre del proyecto",
        value=st.session_state.proyecto_actual
    )

    if st.button("Guardar proyecto"):
        if not nombre_proyecto.strip():
            st.warning("Ingrese un nombre para el proyecto.")
        else:
            datos_guardar = {
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "fuentes_extraidas": st.session_state.fuentes_extraidas,
                "matriz_bibliografica": st.session_state.matriz_bibliografica,
                "referencias_extraidas": st.session_state.referencias_extraidas,
                "articulo_completo": st.session_state.articulo_completo,
                "evaluacion_calidad": st.session_state.evaluacion_calidad,
                "revision_coherencia": st.session_state.revision_coherencia,
                "resumen_estructurado": st.session_state.resumen_estructurado,
                "secciones_generadas": st.session_state.secciones_generadas,
                "datos_articulo": st.session_state.get("datos_articulo", {})
            }

            ruta = guardar_proyecto(nombre_proyecto, datos_guardar)
            st.session_state.proyecto_actual = nombre_proyecto

            st.success(f"Proyecto guardado en: {ruta}")

    st.subheader("Cargar proyecto existente")

    proyectos = listar_proyectos()

    if proyectos:
        proyecto_seleccionado = st.selectbox(
            "Seleccione un proyecto",
            proyectos
        )

        if st.button("Cargar proyecto"):
            ruta = os.path.join("proyectos_guardados", proyecto_seleccionado)
            datos = cargar_proyecto(ruta)

            st.session_state.fuentes_extraidas = datos.get("fuentes_extraidas", "")
            st.session_state.matriz_bibliografica = datos.get("matriz_bibliografica", "")
            st.session_state.referencias_extraidas = datos.get("referencias_extraidas", "")
            st.session_state.articulo_completo = datos.get("articulo_completo", "")
            st.session_state.evaluacion_calidad = datos.get("evaluacion_calidad", "")
            st.session_state.revision_coherencia = datos.get("revision_coherencia", "")
            st.session_state.resumen_estructurado = datos.get("resumen_estructurado", "")
            st.session_state.secciones_generadas = datos.get("secciones_generadas", {})
            st.session_state.datos_articulo = datos.get("datos_articulo", {})
            st.session_state.proyecto_actual = proyecto_seleccionado.replace(".json", "")

            st.success("Proyecto cargado correctamente.")
    else:
        st.info("Aún no hay proyectos guardados.")

    st.subheader("Recomendaciones de uso profesional")

    st.markdown(
        """
        **Flujo recomendado:**

        1. Cargue los artículos PDF.
        2. Extraiga el texto.
        3. Genere la matriz bibliográfica.
        4. Organice las referencias.
        5. Complete título, tema, objetivo y metodología.
        6. Seleccione una plantilla.
        7. Genere el artículo por secciones si desea mayor calidad.
        8. Revise la coherencia interna.
        9. Evalúe la calidad académica.
        10. Genere resumen, abstract y palabras clave.
        11. Exporte en Word para edición final.
        """
    )